#!/usr/bin/env python3
"""
Word文档导出功能
将电子病历导出为Word文档格式
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from typing import Dict, Any, Optional

class WordExporter:
    """Word文档导出器"""
    
    def __init__(self):
        self.document = None
    
    def create_medical_record_document(self, record_data: Dict[str, Any], patient_info: Optional[Dict] = None) -> Document:
        """
        创建电子病历Word文档
        
        Args:
            record_data: 病历数据
            patient_info: 患者信息（可选）
            
        Returns:
            Word文档对象
        """
        # 创建新文档
        self.document = Document()
        
        # 设置页面边距
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 添加标题
        self._add_title()
        
        # 添加基本信息
        if patient_info:
            self._add_patient_info(patient_info)
        
        # 添加病历内容
        self._add_medical_record_content(record_data)
        
        # 添加页脚
        self._add_footer()
        
        return self.document
    
    def _add_title(self):
        """添加文档标题"""
        title = self.document.add_heading('风湿免疫科电子病历', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        time_paragraph = self.document.add_paragraph()
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_run = time_paragraph.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        time_run.font.size = Pt(12)
        time_run.font.color.rgb = None  # 黑色
        
        # 添加分隔线
        self.document.add_paragraph('_' * 50)
    
    def _add_patient_info(self, patient_info: Dict):
        """添加患者基本信息"""
        self.document.add_heading('患者基本信息', level=1)
        
        info_table = self.document.add_table(rows=1, cols=2)
        info_table.style = 'Table Grid'
        
        # 设置表格标题
        header_cells = info_table.rows[0].cells
        header_cells[0].text = '项目'
        header_cells[1].text = '内容'
        
        # 添加患者信息
        info_items = [
            ('姓名', patient_info.get('name', '')),
            ('性别', patient_info.get('gender', '')),
            ('年龄', patient_info.get('age', '')),
            ('住院号', patient_info.get('hospital_id', '')),
            ('科室', '风湿免疫科'),
            ('主治医师', patient_info.get('doctor', '')),
        ]
        
        for item, value in info_items:
            row_cells = info_table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(value)
        
        self.document.add_paragraph()  # 添加空行
    
    def _add_medical_record_content(self, record_data: Dict[str, Any]):
        """添加病历内容"""
        # 定义章节标题和对应的数据键
        sections = [
            ('主诉', 'chief_complaint'),
            ('现病史', 'present_illness'),
            ('既往史', 'past_history'),
            ('体格检查', 'physical_examination'),
            ('辅助检查', 'auxiliary_examination'),
            ('诊断', 'diagnosis'),
            ('治疗方案', 'treatment_plan')
        ]
        
        for section_title, data_key in sections:
            content = record_data.get(data_key, '')
            if content and content.strip() and content.strip() != '无':
                self._add_section(section_title, content)
    
    def _add_section(self, title: str, content: str):
        """添加章节内容"""
        # 添加章节标题
        heading = self.document.add_heading(title, level=1)
        heading.style.font.size = Pt(14)
        heading.style.font.bold = True
        
        # 添加章节内容
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(content)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        
        # 添加空行
        self.document.add_paragraph()
    
    def _add_footer(self):
        """添加页脚"""
        # 添加分隔线
        self.document.add_paragraph('_' * 50)
        
        # 添加签名区域
        signature_para = self.document.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # 添加医师签名
        signature_run = signature_para.add_run('医师签名：_________________')
        signature_run.font.size = Pt(12)
        
        # 添加日期
        date_para = self.document.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f'日期：{datetime.now().strftime("%Y年%m月%d日")}')
        date_run.font.size = Pt(12)
    
    def export_to_bytes(self, record_data: Dict[str, Any], patient_info: Optional[Dict] = None) -> bytes:
        """
        导出为字节数据
        
        Args:
            record_data: 病历数据
            patient_info: 患者信息
            
        Returns:
            Word文档的字节数据
        """
        doc = self.create_medical_record_document(record_data, patient_info)
        
        # 保存到内存
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    def export_to_file(self, record_data: Dict[str, Any], file_path: str, patient_info: Optional[Dict] = None):
        """
        导出到文件
        
        Args:
            record_data: 病历数据
            file_path: 文件路径
            patient_info: 患者信息
        """
        doc = self.create_medical_record_document(record_data, patient_info)
        doc.save(file_path)

def create_word_exporter() -> WordExporter:
    """创建Word导出器实例"""
    return WordExporter() 